"""
Export module for multiple formats (Markdown, PDF, DOCX)
Handles rendering concept note to different file formats
"""

from pathlib import Path
from typing import Dict, Any
from jinja2 import Template
from io import BytesIO
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Note: weasyprint import is deferred to function level to avoid Windows library load issues
PDF_AVAILABLE = True


TEMPLATE_PATH = Path(__file__).parent / "template.md"


def render_markdown(state: Dict[str, Any]) -> str:
    """
    Render the concept note to Markdown using Jinja2 template
    
    Args:
        state: Full wizard state dictionary
        
    Returns:
        Rendered Markdown string
    """
    with open(TEMPLATE_PATH, "r", encoding="utf-8") as f:
        template_content = f.read()
    
    template = Template(template_content)
    markdown = template.render(**state)
    
    return markdown


def markdown_to_html(markdown_content: str) -> str:
    """
    Convert Markdown to HTML for PDF/DOCX export
    
    Args:
        markdown_content: Markdown string
        
    Returns:
        HTML string
    """
    try:
        import markdown
        html = markdown.markdown(markdown_content, extensions=['tables', 'nl2br'])
    except ImportError:
        # Fallback: simple markdown-like conversion
        html = markdown_content.replace('\n\n', '</p><p>')
        html = f"<p>{html}</p>"
    
    return html


def export_to_docx(state: Dict[str, Any]) -> BytesIO:
    """
    Export concept note to DOCX format
    
    Args:
        state: Full wizard state dictionary
        
    Returns:
        BytesIO object containing DOCX file
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed. Install with: pip install python-docx")
    
    # Create a new Document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    if state.get("meta", {}).get("title"):
        title = doc.add_heading(state["meta"]["title"], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    meta_table = doc.add_table(rows=1, cols=2)
    meta_table.style = 'Light Grid Accent 1'
    
    meta_info = [
        ("PIs", ", ".join(state.get("meta", {}).get("pis", []))),
        ("Country", state.get("meta", {}).get("country")),
        ("Funder", state.get("meta", {}).get("funder")),
        ("Stage", state.get("meta", {}).get("project_stage")),
    ]
    
    for label, value in meta_info:
        if value:
            row = meta_table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value)
    
    doc.add_paragraph()  # Spacing
    
    # Problem & Policy Relevance
    doc.add_heading("1. Problem & Policy Relevance", 1)
    
    if state.get("problem_policy_relevance", {}).get("problem_statement"):
        doc.add_heading("Problem Statement", 2)
        doc.add_paragraph(state["problem_policy_relevance"]["problem_statement"])
    
    if state.get("problem_policy_relevance", {}).get("evidence_gap"):
        doc.add_heading("Evidence Gap", 2)
        doc.add_paragraph(state["problem_policy_relevance"]["evidence_gap"])
    
    if state.get("problem_policy_relevance", {}).get("beneficiaries"):
        doc.add_heading("Beneficiaries", 2)
        doc.add_paragraph(state["problem_policy_relevance"]["beneficiaries"])
    
    if state.get("problem_policy_relevance", {}).get("policy_priority"):
        doc.add_heading("Policy Priority", 2)
        doc.add_paragraph(state["problem_policy_relevance"]["policy_priority"])
    
    # Theory of Change
    doc.add_heading("2. Theory of Change", 1)
    
    toc = state.get("theory_of_change", {})
    
    if toc.get("activities"):
        doc.add_heading("Activities", 2)
        for activity in toc["activities"]:
            doc.add_paragraph(activity, style='List Bullet')
    
    if toc.get("outputs"):
        doc.add_heading("Outputs", 2)
        for output in toc["outputs"]:
            doc.add_paragraph(output, style='List Bullet')
    
    if toc.get("short_run_outcomes"):
        doc.add_heading("Short-Run Outcomes", 2)
        for outcome in toc["short_run_outcomes"]:
            doc.add_paragraph(outcome, style='List Bullet')
    
    if toc.get("long_run_outcomes"):
        doc.add_heading("Long-Run Outcomes", 2)
        for outcome in toc["long_run_outcomes"]:
            doc.add_paragraph(outcome, style='List Bullet')
    
    if toc.get("assumptions"):
        doc.add_heading("Key Assumptions", 2)
        for assumption in toc["assumptions"]:
            doc.add_paragraph(assumption, style='List Bullet')
    
    if toc.get("risks"):
        doc.add_heading("Risks", 2)
        for risk in toc["risks"]:
            doc.add_paragraph(risk, style='List Bullet')
    
    # Intervention & Implementation
    doc.add_heading("3. Intervention & Implementation", 1)
    
    intervention = state.get("intervention_implementation", {})
    
    if intervention.get("description"):
        doc.add_paragraph(intervention["description"])
    
    if intervention.get("components"):
        doc.add_heading("Components", 2)
        for component in intervention["components"]:
            doc.add_paragraph(component, style='List Bullet')
    
    if intervention.get("delivery_channels"):
        doc.add_heading("Delivery Channels", 2)
        for channel in intervention["delivery_channels"]:
            doc.add_paragraph(channel, style='List Bullet')
    
    if intervention.get("frequency_intensity"):
        doc.add_heading("Frequency/Intensity", 2)
        doc.add_paragraph(intervention["frequency_intensity"])
    
    if intervention.get("eligibility_criteria"):
        doc.add_heading("Eligibility Criteria", 2)
        doc.add_paragraph(intervention["eligibility_criteria"])
    
    if intervention.get("implementers"):
        doc.add_heading("Implementers", 2)
        for implementer in intervention["implementers"]:
            doc.add_paragraph(implementer, style='List Bullet')
    
    if intervention.get("operational_constraints"):
        doc.add_heading("Operational Constraints", 2)
        for constraint in intervention["operational_constraints"]:
            doc.add_paragraph(constraint, style='List Bullet')
    
    # Study Population & Sampling
    doc.add_heading("4. Study Population & Sampling Frame", 1)
    
    population = state.get("study_population_sampling", {})
    
    if population.get("unit_of_randomization"):
        doc.add_heading("Unit of Randomization", 2)
        doc.add_paragraph(population["unit_of_randomization"])
    
    if population.get("expected_total_n"):
        doc.add_heading("Expected Total N", 2)
        doc.add_paragraph(str(population["expected_total_n"]))
    
    if population.get("sampling_frame_source"):
        doc.add_heading("Sampling Frame Source", 2)
        doc.add_paragraph(population["sampling_frame_source"])
    
    if population.get("inclusion_exclusion"):
        doc.add_heading("Inclusion/Exclusion Criteria", 2)
        doc.add_paragraph(population["inclusion_exclusion"])
    
    if population.get("coverage_limitations"):
        doc.add_heading("Coverage Limitations", 2)
        doc.add_paragraph(population["coverage_limitations"])
    
    # Outcomes & Measurement
    doc.add_heading("5. Outcomes & Measurement", 1)
    
    outcomes = state.get("outcomes_measurement", {})
    
    if outcomes.get("primary_outcomes"):
        doc.add_heading("Primary Outcomes", 2)
        for outcome in outcomes["primary_outcomes"]:
            doc.add_paragraph(outcome, style='List Bullet')
    
    if outcomes.get("secondary_outcomes"):
        doc.add_heading("Secondary Outcomes", 2)
        for outcome in outcomes["secondary_outcomes"]:
            doc.add_paragraph(outcome, style='List Bullet')
    
    if outcomes.get("measurement_timing"):
        timing = outcomes["measurement_timing"]
        if any(timing.values()):
            doc.add_heading("Measurement Timing", 2)
            if timing.get("baseline"):
                doc.add_paragraph(f"Baseline: {timing['baseline']}", style='List Bullet')
            if timing.get("midline"):
                doc.add_paragraph(f"Midline: {timing['midline']}", style='List Bullet')
            if timing.get("endline"):
                doc.add_paragraph(f"Endline: {timing['endline']}", style='List Bullet')
    
    if outcomes.get("instruments"):
        doc.add_heading("Measurement Instruments", 2)
        for instrument in outcomes["instruments"]:
            doc.add_paragraph(instrument, style='List Bullet')
    
    # Randomization Design
    doc.add_heading("6. Randomization Design", 1)
    
    randomization = state.get("randomization_design", {})
    
    numeric = randomization.get("numeric", {})
    if numeric:
        doc.add_heading("Design Parameters", 2)
        if numeric.get("design_type"):
            doc.add_paragraph(f"Design Type: {numeric['design_type']}")
        if numeric.get("arms"):
            doc.add_paragraph(f"Number of Arms: {numeric['arms']}")
        if numeric.get("seed"):
            doc.add_paragraph(f"Seed: {numeric['seed']}")
        if numeric.get("strata"):
            strata_str = ", ".join(numeric["strata"]) if isinstance(numeric["strata"], list) else str(numeric["strata"])
            doc.add_paragraph(f"Strata: {strata_str}")
    
    narrative = randomization.get("narrative", {})
    if narrative.get("rationale"):
        doc.add_heading("Rationale", 2)
        doc.add_paragraph(narrative["rationale"])
    
    if narrative.get("implementation_steps"):
        doc.add_heading("Implementation Steps", 2)
        doc.add_paragraph(narrative["implementation_steps"])
    
    if narrative.get("concealment"):
        doc.add_heading("Concealment Strategy", 2)
        doc.add_paragraph(narrative["concealment"])
    
    # Power & Sample Size
    doc.add_heading("7. Power & Sample Size", 1)
    
    power = state.get("power_sample_size", {})
    
    numeric = power.get("numeric", {})
    if numeric:
        doc.add_heading("Sample Size & Power", 2)
        if numeric.get("n_per_arm"):
            doc.add_paragraph(f"N per Arm: {numeric['n_per_arm']}")
        if numeric.get("mde"):
            doc.add_paragraph(f"MDE: {numeric['mde']}")
        if numeric.get("alpha"):
            doc.add_paragraph(f"Alpha: {numeric['alpha']}")
        if numeric.get("power"):
            doc.add_paragraph(f"Power: {numeric['power']}")
    
    # Data Collection Plan
    doc.add_heading("8. Data Collection Plan", 1)
    
    data_collection = state.get("data_collection_plan", {})
    
    if data_collection.get("mode"):
        doc.add_heading("Data Collection Mode", 2)
        doc.add_paragraph(data_collection["mode"])
    
    if data_collection.get("survey_schedule"):
        doc.add_heading("Survey Schedule", 2)
        doc.add_paragraph(data_collection["survey_schedule"])
    
    if data_collection.get("qc_protocols"):
        doc.add_heading("Quality Control Protocols", 2)
        doc.add_paragraph(data_collection["qc_protocols"])
    
    # Ethics & Approvals
    doc.add_heading("9. Ethics, Risks & Approvals", 1)
    
    ethics = state.get("ethics_risks_approvals", {})
    
    if ethics.get("irb_status"):
        doc.add_heading("IRB Status", 2)
        doc.add_paragraph(ethics["irb_status"])
    
    if ethics.get("consent_process"):
        doc.add_heading("Consent Process", 2)
        doc.add_paragraph(ethics["consent_process"])
    
    if ethics.get("privacy_security"):
        doc.add_heading("Privacy & Security", 2)
        doc.add_paragraph(ethics["privacy_security"])
    
    # Budget Summary
    doc.add_heading("10. Budget Summary", 1)
    
    budget = state.get("budget_summary", {})
    
    if budget.get("categories"):
        categories = budget["categories"]
        doc.add_heading("Budget Breakdown", 2)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        header = table.rows[0]
        header.cells[0].text = "Category"
        header.cells[1].text = "Amount"
        
        total = 0
        for cat, amount in categories.items():
            if amount:
                row = table.add_row()
                row.cells[0].text = cat.replace('_', ' ').title()
                row.cells[1].text = f"${amount:,.2f}"
                total += amount
        
        if total > 0:
            row = table.add_row()
            row.cells[0].text = "Total"
            row.cells[1].text = f"${total:,.2f}"
    
    # Footer
    doc.add_paragraph()
    footer_text = f"Generated by RCT Design Wizard v1.0.0 on {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    footer_paragraph = doc.add_paragraph(footer_text)
    footer_paragraph.runs[0].italic = True
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output


def export_to_pdf(state: Dict[str, Any]) -> BytesIO:
    """
    Export concept note to PDF format
    
    Args:
        state: Full wizard state dictionary
        
    Returns:
        BytesIO object containing PDF file
    """
    # Import at function level to avoid Windows library issues
    try:
        from weasyprint import HTML
    except ImportError as e:
        raise ImportError(f"weasyprint not installed. Install with: pip install weasyprint.\n\nError: {str(e)}")
    
    # Generate markdown first
    markdown_content = render_markdown(state)
    
    # Convert to HTML
    html_content = markdown_to_html(markdown_content)
    
    # Add CSS styling
    css_string = """
    <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            margin: 1in;
        }
        h1 {
            color: #164a7f;
            border-bottom: 3px solid #164a7f;
            padding-bottom: 10px;
            margin-top: 30px;
            page-break-after: avoid;
        }
        h2 {
            color: #2fa6dc;
            margin-top: 20px;
            page-break-after: avoid;
        }
        h3 {
            color: #333;
            margin-top: 15px;
        }
        p {
            margin: 10px 0;
        }
        ul, ol {
            margin: 10px 0 10px 30px;
        }
        li {
            margin: 5px 0;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }
        th, td {
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }
        th {
            background-color: #164a7f;
            color: white;
        }
        tr:nth-child(even) {
            background-color: #f9f9f9;
        }
        .footer {
            text-align: center;
            margin-top: 40px;
            font-size: 10px;
            color: #666;
            border-top: 1px solid #ddd;
            padding-top: 20px;
        }
    </style>
    """
    
    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        {css_string}
    </head>
    <body>
        {html_content}
        <div class="footer">
            Generated by RCT Design Wizard v1.0.0 on {datetime.now().strftime('%Y-%m-%d %H:%M')}
        </div>
    </body>
    </html>
    """
    
    # Convert to PDF
    output = BytesIO()
    HTML(string=full_html).write_pdf(output)
    output.seek(0)
    
    return output


def export_to_file(state: Dict[str, Any], output_path: str, format_type: str = "markdown") -> bool:
    """
    Export concept note to a file
    
    Args:
        state: Full wizard state dictionary
        output_path: Path to save the file
        format_type: Format to export ('markdown', 'docx', 'pdf')
        
    Returns:
        True if successful, False otherwise
    """
    try:
        if format_type.lower() == "markdown":
            markdown = render_markdown(state)
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(markdown)
        elif format_type.lower() == "docx":
            docx_bytes = export_to_docx(state)
            with open(output_path, "wb") as f:
                f.write(docx_bytes.getvalue())
        elif format_type.lower() == "pdf":
            pdf_bytes = export_to_pdf(state)
            with open(output_path, "wb") as f:
                f.write(pdf_bytes.getvalue())
        else:
            raise ValueError(f"Unknown format: {format_type}")
        
        return True
    except Exception as e:
        print(f"Error exporting to {output_path}: {e}")
        return False
